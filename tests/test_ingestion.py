from docx import Document

from ingestion import chunk_text, load_docx, load_pdf, load_txt


def test_txt_returns_text(tmp_path):
    sample = tmp_path / "sample.txt"
    sample.write_text("hello fraud pipeline", encoding="utf-8")
    text = load_txt(str(sample))
    assert isinstance(text, str)
    assert "hello" in text


def test_pdf_returns_text():
    text = load_pdf("/home/ubuntu/Desktop/RAG Architecture.pdf")
    assert len(text) > 100


def test_chunks_are_correct_size():
    chunks = chunk_text("word " * 2000)
    assert all(len(c) <= 600 for c in chunks)


def test_docx_returns_text(tmp_path):
    sample = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Sample docx content for ingestion tests.")
    doc.save(sample)
    text = load_docx(str(sample))
    assert len(text) > 0
    assert "Sample docx" in text


def test_docx_returns_string(tmp_path):
    sample = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Another paragraph.")
    doc.save(sample)
    text = load_docx(str(sample))
    assert isinstance(text, str)
